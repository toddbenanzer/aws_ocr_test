from docx import Document
from docx.shared import Inches
import random
import lorem
import textgen


class DocGen(object):
    def __init__(self):
        self.doc = Document()

    def gen_heading(self):
        return textgen.generate_heading(length=10, base_test=None)
        #return lorem.text()[:random.randint(8, 20)]

    def gen_paragraph(self):
        return textgen.generate_paragraph(length=600, base_test=None)
        #return lorem.paragraph()

    def add_heading(self, text='', level=1):
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text=''):
        self.doc.add_paragraph(text)

    def add_page_break(self):
        self.doc.add_page_break()

    def add_section(self, num_paragraphs, level=1):
        heading = self.gen_heading()
        self.doc.add_heading(heading, level=level)
        for i in range(num_paragraphs):
            paragraph = self.gen_paragraph()
            self.add_paragraph(paragraph)

    def save(self, filename):
        self.doc.save(filename)


def generate_random_document(filename, num_sections=random.randint(10, 40)):
    doc = DocGen()
    curr_level = 1
    for section in range(num_sections):
        curr_level = random.choice([1, curr_level, curr_level + 1])
        num_paragraphs = random.randint(1, 10)
        doc.add_section(num_paragraphs, level=curr_level)
        if random.random() > 0.75:
            doc.add_page_break()
    doc.save(filename)


generate_random_document(r'C:\developer\aws_ocr_test\created_docs\test_gpt2.docx')
